"""Word（.docx）解析策略。"""

from __future__ import annotations

from pathlib import Path

from docx import Document as DocxDocument

from app.services.document_parser.base import DocumentParseError, DocumentParseStrategy


class DocxParseStrategy(DocumentParseStrategy):
    """提取 Word 段落文本为全文（不含页眉页脚/复杂表格布局）。"""

    def extract_text(self, file_path: Path) -> str:
        """从 docx 提取段落文本。

        Args:
            file_path: 本地 docx 文件路径。

        Returns:
            段落用换行拼接后的全文。

        Raises:
            DocumentParseError: 文件损坏或无法读取时抛出。
        """
        if not file_path.is_file():
            raise DocumentParseError("Word 文件不存在")
        try:
            document = DocxDocument(str(file_path))
            parts: list[str] = []
            for paragraph in document.paragraphs:
                text = (paragraph.text or "").strip()
                if text:
                    parts.append(text)
            # 简单表格：按单元格空格拼接，行之间换行
            for table in document.tables:
                for row in table.rows:
                    cells = [
                        (cell.text or "").strip()
                        for cell in row.cells
                        if (cell.text or "").strip()
                    ]
                    if cells:
                        parts.append(" ".join(cells))
            return "\n".join(parts).strip()
        except DocumentParseError:
            raise
        except Exception as exc:  # noqa: BLE001
            raise DocumentParseError(f"Word 解析失败: {exc}") from exc
